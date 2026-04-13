from docx import Document

def create_summary_docx():
    doc = Document()
    doc.add_heading('CS229 Course Book Summaries', 0)

    summaries = [
        ("Chapter 1: Linear Regression", 
         "Introduces supervised learning with continuous target variables. It covers the LMS (Least Mean Squares) algorithm, gradient descent, and the cost function J(θ) used to minimize prediction error."),
        
        ("Chapter 2: Classification and Logistic Regression", 
         "Focuses on classification problems where target variables are discrete. It introduces logistic regression, the perceptron learning algorithm, and locally weighted linear regression."),
        
        ("Chapter 3: Generalized Linear Models (GLM)", 
         "Explains the construction of GLMs using the exponential family of distributions. It shows how ordinary least squares and logistic regression are special cases of this broader framework."),
        
        ("Chapter 4: Generative Learning Algorithms", 
         "Covers algorithms like Gaussian Discriminant Analysis (GDA) and Naive Bayes. It compares generative models (which model p(x|y)) with discriminative models (which model p(y|x))."),
        
        ("Chapter 5: Kernel Methods", 
         "Introduces feature maps and the 'kernel trick.' This allows algorithms to efficiently operate in high-dimensional feature spaces by computing inner products without explicitly mapping features."),
        
        ("Chapter 6: Support Vector Machines (SVM)", 
         "Discusses optimal margin classifiers, Lagrange duality, and the application of kernels to SVMs for robust non-linear classification."),
        
        ("Chapter 7: Deep Learning", 
         "Provides an overview of neural networks and non-linear models. It covers vectorization and the backpropagation algorithm used for training complex models."),
        
        ("Chapter 8: Generalization", 
         "Analyzes how models perform on unseen test data. It discusses the relationship between training loss and test error, as well as the bias-variance tradeoff."),
        
        ("Chapter 9: Regularization and Model Selection", 
         "Explores techniques to prevent overfitting, including L1/L2 regularization, cross-validation for model selection, and Bayesian statistical perspectives."),
        
        ("Chapter 10: Clustering and the K-Means Algorithm", 
         "Introduces unsupervised learning through clustering. It details the K-means algorithm for grouping unlabeled data into cohesive clusters."),
        
        ("Chapter 11: EM Algorithms", 
         "Discusses the Expectation-Maximization (EM) algorithm, particularly for Mixture of Gaussians. It explains how to handle latent variables in probabilistic models."),
        
        ("Chapter 12: Principal Components Analysis (PCA)", 
         "Covers dimensionality reduction by identifying the principal components that preserve the most variance in a dataset, useful for compression and visualization."),
        
        ("Chapter 13: Independent Components Analysis (ICA)", 
         "Focuses on source separation, where the goal is to recover independent underlying signals (like individual voices) from a mixture of observations."),
        
        ("Chapter 14: Self-Supervised Learning and Foundation Models", 
         "Discusses modern techniques like pretraining and adaptation. It covers Transformers, Large Language Models (LLMs), and zero-shot/in-context learning."),
        
        ("Chapter 15: Reinforcement Learning", 
         "Introduces Markov Decision Processes (MDPs). It explains value iteration, policy iteration, and how agents learn to make sequences of decisions to maximize rewards."),
        
        ("Chapter 16: LQR, DDP and LQG", 
         "Discusses control in finite-horizon MDPs, specifically Linear Quadratic Regulation (LQR) and Differential Dynamic Programming (DDP) for non-linear systems."),
        
        ("Chapter 17: Policy Gradient (REINFORCE)", 
         "Explains how to optimize policies directly using gradient ascent. It addresses the challenges of estimating gradients when reward functions or transition dynamics are unknown.")
    ]

    for title, summary in summaries:
        doc.add_heading(title, level=1)
        doc.add_paragraph(summary)

    doc.save('CS229_Course_Summaries.docx')
    print("Document 'CS229_Course_Summaries.docx' created successfully.")

if __name__ == "__main__":
    create_summary_docx()
